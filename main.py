from docx import Document
from bs4 import BeautifulSoup


def create_soup_from_html_file(html_file_name):
    """
    Creates a BeautifulSoup object with the data from the HTML file you pass it

    :param html_file_name: the exact name of the HTML file as it appears in the "Downloads" folder
    :return: a BeautifulSoup object
    """
    full_file_path = f"C:/Users/19493/Downloads/{html_file_name}"
    html_content = open(full_file_path, "r", encoding="utf-8").read()
    new_soup = BeautifulSoup(html_content, "html.parser")
    return new_soup


def create_heading(chosen_document, heading_text):
    chosen_document.add_heading(heading_text)


def create_points_from_html_element(the_div):
    """
    Takes in a html element (most likely a div) and parses through the element to find children elements we can
    add to a list it will then return. This function gets called by create_bulleted_list because it needs
    a list of bullet points.

    :param the_div: the div element that contains bullet points (li's)
    :return: a list of points that we want to add to the document
    """
    points_to_add = []
    div_children = the_div.contents
    for ele in div_children:
        # There still might be \n in the middle of a bullet point because of weird format conversion
        # from a ppt to html file, so we have to take that into account
        eles_split_list = ele.text.split("\n")
        for each_ele in eles_split_list:
            # Now keep out blank space and /n elements
            if each_ele != "\n" and each_ele:
                points_to_add.append(each_ele)
    return points_to_add


def create_bulleted_list(chosen_document, the_element):
    """
    Takes in an element that contains points (child elements) you want to add, and creates a bullet point list on
    the Word document.

    :param chosen_document: The Word document you are creating
    :param the_element: The element that has the points you want to add as children elements
    """
    list_of_bullet_points = create_points_from_html_element(the_element)
    for point in list_of_bullet_points:
        chosen_document.add_paragraph(point, style='List Bullet')


def make_img_readable(src):
    pass


# Create the soup object and Word doc object
soup = create_soup_from_html_file("chap01.html")
document = Document()

element_list = soup.body.contents

# For each element that is enclosed in the body tag
for element in element_list:
    if element.name == "h1":    # h1 elements are section titles for the chapter
        create_heading(document, element.text)
        print(element.text)

    elif element.name == "div":     # divs store bullet points with info on each section
        create_bulleted_list(document, element)

    elif element.name == "figure":     # graphs/figures/etc. will need to be inserted as well
        img_src = element.find('img')['src']
        document.add_picture(img_src)

    else:
        # Then it is none of these, then just put it in as plain text because we can't identify it
        pass


document.save("C:/Users/19493/Desktop/new_word_doc.docx")
